"""
Convierte los manuales de FAControl (.md) a Word (.docx), metiendo las
capturas de pantalla donde estan los marcadores de imagen.

POR QUE EXISTE
--------------
Los manuales se escriben en Markdown porque es texto plano: se versiona, se
diffea y se corrige rapido. Pero al cliente hay que entregarle un .docx que
pueda abrir, imprimir y anotar. Este script hace ese ultimo paso, y sobre todo
resuelve el trabajo aburrido: pegar 60 capturas en el lugar exacto sin
equivocarse.

COMO SE USA
-----------
    python scripts/docs/generar_docx.py

    # con las capturas en otra carpeta
    python scripts/docs/generar_docx.py --imagenes "C:\\ruta\\a\\capturas"

DONDE VAN LAS CAPTURAS
----------------------
En `docs/imagenes/`, con el nombre que dice el marcador:

    imagen-01.png ... imagen-46.png     -> las del MANUAL
    imagen-I-01.png ... imagen-I-14.png -> las de la GUIA DE INSTALACION

Sirven .png, .jpg y .jpeg: el script prueba las tres extensiones.

La captura que FALTA no rompe nada: en su lugar queda un recuadro gris con el
texto de lo que habria que fotografiar. Asi se puede generar el .docx a mitad
de camino y ver como va quedando.

REQUIERE
--------
    pip install python-docx
"""

from __future__ import annotations

import argparse
import io
import os
import re
import sys

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    sys.exit("Falta python-docx. Instalalo con:  pip install python-docx")


RAIZ = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DOCS = os.path.join(RAIZ, "docs")

# Documentos a convertir: (archivo .md, archivo .docx, titulo de portada)
DOCUMENTOS = [
    ("MANUAL.md", "FAControl - Manual de usuario.docx", "Manual de usuario"),
    ("INSTALL.md", "FAControl - Guia de instalacion.docx", "Guía de instalación"),
    # No lleva capturas: son pantallas de Google, que cambian solas cada tanto y
    # dejarian el documento desactualizado sin que nadie se entere.
    ("CORREO-GMAIL.md", "FAControl - Correo automatico.docx", "Correo automático"),
]

ANCHO_IMAGEN = Inches(6.0)   # ancho util de una hoja carta con margenes de 1"

# El marcador de imagen ocupa 3 lineas dentro de un bloque ``` y se ve asi:
#     ┌───────────────┐
#     │  📷 IMAGEN 07 — Descripcion, que puede seguir en la linea de abajo  │
#     └───────────────┘
RE_MARCADOR = re.compile(r"📷\s*IMAGEN\s+([0-9I\-]+)\s*[—–-]\s*(.*)", re.IGNORECASE)


# ---------------------------------------------------------------------------
# Utilidades de formato
# ---------------------------------------------------------------------------

def sombrear(celda, hex_color: str) -> None:
    """Color de fondo de una celda (python-docx no lo expone)."""
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:val"), "clear")
    sombra.set(qn("w:fill"), hex_color)
    celda._tc.get_or_add_tcPr().append(sombra)


def borde_parrafo(parrafo, hex_color: str = "BBBBBB") -> None:
    """Recuadro alrededor de un parrafo: lo usa el hueco de imagen faltante."""
    pPr = parrafo._p.get_or_add_pPr()
    bordes = OxmlElement("w:pBdr")
    for lado in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{lado}")
        b.set(qn("w:val"), "dashed")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:space"), "6")
        b.set(qn("w:color"), hex_color)
        bordes.append(b)
    pPr.append(bordes)


def escribir_con_formato(parrafo, texto: str) -> None:
    """
    Escribe el texto respetando **negrita**, *cursiva* y `codigo`.

    Un parser de Markdown completo seria exagerado: estos manuales usan tres
    marcas y nada mas. Se recorren con una sola expresion para no perder el
    orden en que aparecen.
    """
    partes = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", texto)
    for parte in partes:
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**") and len(parte) > 4:
            parrafo.add_run(parte[2:-2]).bold = True
        elif parte.startswith("`") and parte.endswith("`") and len(parte) > 2:
            run = parrafo.add_run(parte[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif parte.startswith("*") and parte.endswith("*") and len(parte) > 2:
            parrafo.add_run(parte[1:-1]).italic = True
        else:
            parrafo.add_run(parte)


def limpiar_enlaces(texto: str) -> str:
    """[texto](destino) -> texto. Los enlaces del indice no sirven en Word."""
    return re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", texto)


# ---------------------------------------------------------------------------
# Imagenes
# ---------------------------------------------------------------------------

def buscar_imagen(carpeta: str, numero: str) -> str | None:
    """La captura de ese numero, probando las extensiones habituales."""
    for ext in (".png", ".jpg", ".jpeg", ".PNG", ".JPG"):
        ruta = os.path.join(carpeta, f"imagen-{numero}{ext}")
        if os.path.isfile(ruta):
            return ruta
    return None


def insertar_imagen(doc, carpeta: str, numero: str, descripcion: str,
                    faltantes: list[str]) -> None:
    """La captura si existe; si no, un hueco marcado para no perderlo de vista."""
    ruta = buscar_imagen(carpeta, numero)

    if ruta:
        parrafo = doc.add_paragraph()
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            parrafo.add_run().add_picture(ruta, width=ANCHO_IMAGEN)
        except Exception as ex:                      # imagen corrupta o vacia
            faltantes.append(f"{numero} (no se pudo leer: {ex})")
            parrafo.add_run(f"[No se pudo insertar imagen-{numero}]")
    else:
        faltantes.append(numero)
        hueco = doc.add_paragraph()
        hueco.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hueco.add_run(f"\n📷  FALTA LA IMAGEN {numero}\n{descripcion}\n")
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        borde_parrafo(hueco)

    # Pie de foto: queda igual esté o no la captura, para poder referenciarla
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pie.add_run(f"Imagen {numero} — {descripcion}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ---------------------------------------------------------------------------
# Conversion
# ---------------------------------------------------------------------------

def leer_marcador(bloque: list[str]) -> tuple[str, str] | None:
    """
    Numero y descripcion de un bloque de codigo que sea un marcador de imagen.

    La descripcion puede venir partida en varias lineas del recuadro, asi que
    se juntan todas las que estan entre las barras.
    """
    texto = " ".join(bloque)
    if "📷" not in texto:
        return None

    coincidencia = RE_MARCADOR.search(texto)
    if not coincidencia:
        return None

    numero = coincidencia.group(1).strip()
    # Se sacan los caracteres del recuadro y los espacios que sobran
    descripcion = coincidencia.group(2)
    descripcion = descripcion.replace("│", " ").replace("└", " ").replace("┘", " ")
    descripcion = descripcion.replace("─", " ").replace("┌", " ").replace("┐", " ")
    return numero, " ".join(descripcion.split())


def agregar_tabla(doc, filas: list[str]) -> None:
    """Una tabla de Markdown (| a | b |) como tabla de Word."""
    datos = []
    for fila in filas:
        celdas = [c.strip() for c in fila.strip().strip("|").split("|")]
        # La linea de guiones (|---|---|) es separador, no datos
        if all(re.fullmatch(r":?-{2,}:?", c) for c in celdas if c):
            continue
        datos.append(celdas)

    if not datos:
        return

    columnas = max(len(f) for f in datos)
    tabla = doc.add_table(rows=0, cols=columnas)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, fila in enumerate(datos):
        celdas_word = tabla.add_row().cells
        for j in range(columnas):
            texto = fila[j] if j < len(fila) else ""
            parrafo = celdas_word[j].paragraphs[0]
            escribir_con_formato(parrafo, limpiar_enlaces(texto))
            if i == 0:
                sombrear(celdas_word[j], "EEF0FE")
                for run in parrafo.runs:
                    run.bold = True
    doc.add_paragraph()


def convertir(ruta_md: str, ruta_docx: str, subtitulo: str,
              carpeta_imagenes: str) -> list[str]:
    """Genera el .docx. Devuelve la lista de imagenes que faltaron."""
    with io.open(ruta_md, encoding="utf-8") as f:
        lineas = f.read().replace("\r\n", "\n").split("\n")

    doc = Document()
    for seccion in doc.sections:
        seccion.left_margin = seccion.right_margin = Inches(1)

    # --- Portada ---
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("\n\n\nFAControl")
    run.bold = True
    run.font.size = Pt(40)
    run.font.color.rgb = RGBColor(0x2B, 0x3A, 0x67)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitulo)
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pie.add_run("\n\nFamilia Almonte Auto Import SRL\nYuber Santana · 849-438-0242")
    run.font.size = Pt(12)

    doc.add_section(WD_SECTION.NEW_PAGE)

    faltantes: list[str] = []
    i = 0
    en_cita = False

    while i < len(lineas):
        linea = lineas[i]
        crudo = linea.rstrip()
        texto = crudo.strip()

        # --- Bloque de codigo (puede ser un marcador de imagen) ---
        if texto.startswith("```"):
            bloque, i = [], i + 1
            while i < len(lineas) and not lineas[i].strip().startswith("```"):
                bloque.append(lineas[i])
                i += 1
            i += 1

            marcador = leer_marcador(bloque)
            if marcador:
                insertar_imagen(doc, carpeta_imagenes, marcador[0], marcador[1], faltantes)
            else:
                parrafo = doc.add_paragraph()
                run = parrafo.add_run("\n".join(l.rstrip() for l in bloque))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                sombrear_parrafo = parrafo._p.get_or_add_pPr()
                sombra = OxmlElement("w:shd")
                sombra.set(qn("w:val"), "clear")
                sombra.set(qn("w:fill"), "F5F5F5")
                sombrear_parrafo.append(sombra)
            continue

        # --- Tabla ---
        if texto.startswith("|"):
            filas = []
            while i < len(lineas) and lineas[i].strip().startswith("|"):
                filas.append(lineas[i])
                i += 1
            agregar_tabla(doc, filas)
            continue

        # --- Separador ---
        if texto in ("---", "***", "___"):
            doc.add_paragraph("─" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # --- Linea en blanco ---
        if not texto:
            en_cita = False
            i += 1
            continue

        # --- Titulos ---
        encabezado = re.match(r"^(#{1,4})\s+(.*)", texto)
        if encabezado:
            nivel = len(encabezado.group(1))
            doc.add_heading(limpiar_enlaces(encabezado.group(2)), level=nivel)
            i += 1
            continue

        # --- Cita / recuadro de aviso ---
        if texto.startswith(">"):
            contenido = texto.lstrip(">").strip()
            if not contenido:
                i += 1
                continue
            # Un titulo dentro de una cita sigue siendo un titulo
            enc_cita = re.match(r"^(#{1,4})\s+(.*)", contenido)
            parrafo = doc.add_paragraph()
            parrafo.paragraph_format.left_indent = Inches(0.35)
            if enc_cita:
                run = parrafo.add_run(enc_cita.group(2))
                run.bold = True
                run.font.size = Pt(13)
            else:
                escribir_con_formato(parrafo, limpiar_enlaces(contenido))
                for run in parrafo.runs:
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            borde_parrafo(parrafo, "2B3A67") if not en_cita else None
            en_cita = True
            i += 1
            continue

        # --- Lista con vinetas ---
        vineta = re.match(r"^[-*]\s+(.*)", texto)
        if vineta:
            contenido = vineta.group(1)
            # Casillas de verificacion: [ ] y [x]
            casilla = re.match(r"^\[([ xX])\]\s*(.*)", contenido)
            parrafo = doc.add_paragraph(style="List Bullet")
            if casilla:
                marca = "☑ " if casilla.group(1).lower() == "x" else "☐ "
                parrafo.add_run(marca)
                contenido = casilla.group(2)
            escribir_con_formato(parrafo, limpiar_enlaces(contenido))
            i += 1
            continue

        # --- Lista numerada ---
        numerada = re.match(r"^\d+\.\s+(.*)", texto)
        if numerada:
            parrafo = doc.add_paragraph(style="List Number")
            escribir_con_formato(parrafo, limpiar_enlaces(numerada.group(1)))
            i += 1
            continue

        # --- Parrafo normal: se juntan las lineas seguidas ---
        partes = [texto]
        i += 1
        while i < len(lineas):
            siguiente = lineas[i].strip()
            if (not siguiente or siguiente.startswith(("#", ">", "|", "```", "-", "*", "---"))
                    or re.match(r"^\d+\.\s", siguiente)):
                break
            partes.append(siguiente)
            i += 1
        parrafo = doc.add_paragraph()
        escribir_con_formato(parrafo, limpiar_enlaces(" ".join(partes)))

    doc.save(ruta_docx)
    return faltantes


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Genera los manuales de FAControl en Word (.docx)")
    parser.add_argument("--imagenes", default=os.path.join(DOCS, "imagenes"),
                        help="Carpeta con las capturas (por defecto docs/imagenes)")
    parser.add_argument("--salida", default=DOCS,
                        help="Carpeta donde dejar los .docx (por defecto docs/)")
    args = parser.parse_args()

    if not os.path.isdir(args.imagenes):
        os.makedirs(args.imagenes, exist_ok=True)
        print(f"Se creo la carpeta de capturas: {args.imagenes}")

    os.makedirs(args.salida, exist_ok=True)
    print(f"Capturas: {args.imagenes}\n")

    for nombre_md, nombre_docx, subtitulo in DOCUMENTOS:
        ruta_md = os.path.join(DOCS, nombre_md)
        if not os.path.isfile(ruta_md):
            print(f"  [!] No existe {nombre_md}, se saltea")
            continue

        ruta_docx = os.path.join(args.salida, nombre_docx)
        faltantes = convertir(ruta_md, ruta_docx, subtitulo, args.imagenes)

        print(f"  {nombre_docx}")
        if faltantes:
            print(f"     faltan {len(faltantes)} captura(s): {', '.join(faltantes)}")
        else:
            print("     todas las capturas estan puestas")

    print("\nListo.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
